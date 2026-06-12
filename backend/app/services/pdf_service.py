"""
Document Processing Service for Purchase Order Extraction
Handles PDF (native + OCR) and Word document (.doc, .docx) text extraction.
"""

import logging
import os
from pathlib import Path
from typing import Any, Iterable, List, Optional, Tuple

from app.services.import_service import MAX_CONSECUTIVE_BLANK_ROWS, MAX_IMPORT_COLUMNS, MAX_SCANNED_ROWS

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = ['.pdf', '.doc', '.docx', '.xlsx', '.xls']


class DocumentExtractionResult:
    def __init__(
        self, text: str, is_ocr: bool = False, page_count: int = 0, confidence: str = "high", file_type: str = "pdf"
    ):
        self.text = text
        self.is_ocr = is_ocr
        self.page_count = page_count
        self.confidence = confidence  # high, medium, low
        self.file_type = file_type  # pdf, docx, doc


# Alias for backward compatibility
PDFExtractionResult = DocumentExtractionResult


def extract_text_from_document(file_path: str) -> DocumentExtractionResult:
    """
    Extract text from PDF or Word document.
    Automatically detects file type and uses appropriate extraction method.

    ``file_path`` may be a local path or an ``s3://`` storage ref; remote refs are
    materialized to a real local file first (pdf2image/pytesseract/antiword need one).
    """
    from app.services.storage_service import is_s3_ref, ref_as_local_path

    if is_s3_ref(file_path):
        with ref_as_local_path(file_path) as local_path:
            return extract_text_from_document(str(local_path))

    path = Path(file_path)
    ext = path.suffix.lower()

    logger.info(f"[DOC_EXTRACT] Processing file: {file_path}")
    logger.info(f"[DOC_EXTRACT] Detected extension: '{ext}'")

    if ext == '.pdf':
        logger.info("[DOC_EXTRACT] Routing to PDF extractor")
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        logger.info(f"[DOC_EXTRACT] Routing to Word extractor for {ext}")
        return extract_text_from_word(file_path)
    elif ext in ['.xlsx', '.xls']:
        logger.info(f"[DOC_EXTRACT] Routing to Excel extractor for {ext}")
        return extract_text_from_excel(file_path)
    else:
        logger.warning(f"[DOC_EXTRACT] Unknown extension: {ext}")
        return DocumentExtractionResult(text="", confidence="low", file_type="unknown")


def extract_text_from_word(doc_path: str) -> DocumentExtractionResult:
    """
    Extract text from Word document (.docx or .doc).
    """
    path = Path(doc_path)
    ext = path.suffix.lower()

    logger.info(f"[WORD] extract_text_from_word called with: {doc_path}")
    logger.info(f"[WORD] Extension detected: '{ext}'")

    if ext == '.docx':
        logger.info("[WORD] Calling _extract_docx_text")
        return _extract_docx_text(doc_path)
    elif ext == '.doc':
        logger.info("[WORD] Calling _extract_doc_text")
        return _extract_doc_text(doc_path)
    else:
        logger.warning(f"[WORD] Unsupported Word extension: {ext}")
        return DocumentExtractionResult(text="", confidence="low", file_type="unknown")


def extract_text_from_excel(excel_path: str) -> DocumentExtractionResult:
    """
    Extract text from Excel document (.xlsx or .xls).

    The scan is bounded with the Import Center's shared caps
    (``app.services.import_service``) so a workbook with a bloated used range —
    one stray formatted cell at XFD1048576 declares the full
    16,384 x 1,048,576 grid, minutes of CPU for a KB-sized file — extracts in
    milliseconds: at most ``MAX_IMPORT_COLUMNS`` columns are read per row, a
    per-sheet run of more than ``MAX_CONSECUTIVE_BLANK_ROWS`` blank rows ends
    that sheet (later sheets are still read), and a workbook-wide cap of
    ``MAX_SCANNED_ROWS`` raw rows stops extraction entirely, returning the text
    gathered so far at "medium" confidence. Text extraction degrades
    gracefully — this function never raises; failures return an empty
    low-confidence result.
    """
    path = Path(excel_path)
    ext = path.suffix.lower()

    logger.info(f"[EXCEL] extract_text_from_excel called with: {excel_path}")
    logger.info(f"[EXCEL] Extension detected: '{ext}'")

    if not os.path.exists(excel_path):
        logger.error(f"[EXCEL] File does not exist: {excel_path}")
        return DocumentExtractionResult(text="", confidence="low", file_type=ext.strip('.'))

    file_size = os.path.getsize(excel_path)
    logger.info(f"[EXCEL] File size: {file_size} bytes")
    if file_size == 0:
        logger.error("[EXCEL] File is empty (0 bytes)")
        return DocumentExtractionResult(text="", confidence="low", file_type=ext.strip('.'))

    all_text: List[str] = []
    scanned_rows = 0

    def _append_sheet(sheet_name: str, sheet_rows: Iterable[Any]) -> bool:
        """Fold one sheet into ``all_text``; False when the workbook-wide cap was hit."""
        nonlocal scanned_rows
        all_text.append(f"--- Sheet: {sheet_name} ---")
        consecutive_blank_rows = 0
        for raw_row in sheet_rows:
            scanned_rows += 1
            if scanned_rows > MAX_SCANNED_ROWS:
                logger.warning(f"[EXCEL] Scanned-row cap hit ({MAX_SCANNED_ROWS:,} raw rows) — returning partial text")
                return False
            cells = [str(cell).strip() for cell in raw_row if cell is not None and str(cell).strip()]
            if not cells:
                consecutive_blank_rows += 1
                if consecutive_blank_rows > MAX_CONSECUTIVE_BLANK_ROWS:
                    return True  # used-range bloat on this sheet — move on to the next sheet
                continue
            consecutive_blank_rows = 0
            all_text.append(" | ".join(cells))
        return True

    try:
        if ext == ".xlsx":
            try:
                from openpyxl import load_workbook
            except ImportError as e:
                logger.error(f"[EXCEL] openpyxl not installed: {e}")
                return DocumentExtractionResult(text="", confidence="low", file_type=ext.strip('.'))
            workbook = load_workbook(excel_path, read_only=True, data_only=True)
            try:
                for ws in workbook.worksheets:
                    if not _append_sheet(ws.title, ws.iter_rows(values_only=True, max_col=MAX_IMPORT_COLUMNS)):
                        break
            finally:
                workbook.close()
        elif ext == ".xls":
            try:
                import xlrd
            except ImportError as e:
                logger.error(f"[EXCEL] xlrd not installed: {e}")
                return DocumentExtractionResult(text="", confidence="low", file_type=ext.strip('.'))
            book = xlrd.open_workbook(excel_path)
            for sheet in book.sheets():
                sheet_rows = (
                    [sheet.cell_value(r, c) for c in range(min(sheet.ncols, MAX_IMPORT_COLUMNS))]
                    for r in range(sheet.nrows)
                )
                if not _append_sheet(sheet.name, sheet_rows):
                    break
        text = "\n".join([line for line in all_text if line.strip()])
        return DocumentExtractionResult(text=text, confidence="medium", file_type=ext.strip('.'))
    except Exception as e:
        logger.error(f"[EXCEL] Extraction failed: {e}")
        return DocumentExtractionResult(text="", confidence="low", file_type=ext.strip('.'))


def _extract_docx_text(docx_path: str) -> DocumentExtractionResult:
    """Extract text from .docx file using python-docx."""
    logger.info(f"[DOCX] Starting extraction from: {docx_path}")

    # Verify file exists
    if not os.path.exists(docx_path):
        logger.error(f"[DOCX] File does not exist: {docx_path}")
        return DocumentExtractionResult(text="", confidence="low", file_type="docx")

    file_size = os.path.getsize(docx_path)
    logger.info(f"[DOCX] File size: {file_size} bytes")

    if file_size == 0:
        logger.error("[DOCX] File is empty (0 bytes)")
        return DocumentExtractionResult(text="", confidence="low", file_type="docx")

    try:
        from docx import Document

        logger.info("[DOCX] python-docx imported successfully")
    except ImportError as e:
        logger.error(f"[DOCX] python-docx not installed: {e}")
        return DocumentExtractionResult(text="", confidence="low", file_type="docx")

    try:
        # Open the document - use absolute path to avoid issues
        abs_path = os.path.abspath(docx_path)
        logger.info(f"[DOCX] Opening document at absolute path: {abs_path}")

        doc = Document(abs_path)
        logger.info("[DOCX] Document opened successfully")

        all_text = []

        # Extract paragraphs
        para_count = len(doc.paragraphs)
        logger.info(f"[DOCX] Found {para_count} paragraphs")

        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())

        logger.info(f"[DOCX] Extracted {len(all_text)} non-empty paragraphs")

        # Extract tables (critical for PO line items)
        table_count = len(doc.tables)
        logger.info(f"[DOCX] Found {table_count} tables")

        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    all_text.append("\t".join(row_text))
            logger.info(f"[DOCX] Table {table_idx + 1}: extracted {len(table.rows)} rows")

        text = "\n".join(all_text)

        logger.info(f"[DOCX] Extraction complete: {len(text)} chars total")

        if not text.strip():
            logger.warning("[DOCX] Document appears to be empty (no text content)")
            return DocumentExtractionResult(text="", confidence="low", file_type="docx")

        return DocumentExtractionResult(text=text, is_ocr=False, page_count=1, confidence="high", file_type="docx")

    except Exception as e:
        logger.error(f"[DOCX] Extraction failed with exception: {type(e).__name__}: {e}")
        import traceback

        logger.error(f"[DOCX] Traceback: {traceback.format_exc()}")
        return DocumentExtractionResult(text="", confidence="low", file_type="docx")


def _extract_doc_text(doc_path: str) -> DocumentExtractionResult:
    """
    Extract text from legacy .doc file.
    Tries multiple methods: antiword, catdoc, or python-docx fallback.
    """
    import subprocess

    logger.info(f"[DOC] Starting legacy .doc extraction from: {doc_path}")

    # Verify file exists and get resolved path (prevents path traversal)
    from pathlib import Path

    abs_path = str(Path(doc_path).resolve())
    # Ensure the resolved path is within the expected upload directory
    allowed_dirs = ["/app/uploads", "/tmp"]  # nosec B108 - allowlist for a path-traversal guard, not temp-file creation
    if not any(abs_path.startswith(d) for d in allowed_dirs):
        logger.error(f"[DOC] Path traversal attempt blocked: {abs_path}")
        return DocumentExtractionResult(text="", confidence="low", file_type="doc")
    if not os.path.exists(abs_path):
        logger.error(f"[DOC] File does not exist: {abs_path}")
        return DocumentExtractionResult(text="", confidence="low", file_type="doc")

    file_size = os.path.getsize(abs_path)
    logger.info(f"[DOC] File size: {file_size} bytes, absolute path: {abs_path}")

    # Method 1: Try antiword (Linux/Mac)
    logger.info("[DOC] Trying antiword...")
    try:
        result = subprocess.run(['antiword', abs_path], capture_output=True, text=True, timeout=30)
        logger.info(
            f"[DOC] antiword returncode: {result.returncode}, stdout length: {len(result.stdout)}, stderr: {result.stderr[:200] if result.stderr else 'none'}"
        )
        if result.returncode == 0 and result.stdout.strip():
            logger.info(f"[DOC] antiword extraction successful: {len(result.stdout)} chars")
            return DocumentExtractionResult(
                text=result.stdout, is_ocr=False, page_count=1, confidence="high", file_type="doc"
            )
        else:
            logger.warning(f"[DOC] antiword returned empty or failed: rc={result.returncode}")
    except FileNotFoundError:
        logger.warning("[DOC] antiword not installed")
    except subprocess.TimeoutExpired:
        logger.warning("[DOC] antiword timed out")
    except Exception as e:
        logger.warning(f"[DOC] antiword exception: {type(e).__name__}: {e}")

    # Method 2: Try catdoc (alternative for .doc)
    logger.info("[DOC] Trying catdoc...")
    try:
        result = subprocess.run(['catdoc', abs_path], capture_output=True, text=True, timeout=30)
        logger.info(f"[DOC] catdoc returncode: {result.returncode}, stdout length: {len(result.stdout)}")
        if result.returncode == 0 and result.stdout.strip():
            logger.info(f"[DOC] catdoc extraction successful: {len(result.stdout)} chars")
            return DocumentExtractionResult(
                text=result.stdout, is_ocr=False, page_count=1, confidence="medium", file_type="doc"
            )
    except FileNotFoundError:
        logger.warning("[DOC] catdoc not installed")
    except subprocess.TimeoutExpired:
        logger.warning("[DOC] catdoc timed out")
    except Exception as e:
        logger.warning(f"[DOC] catdoc exception: {type(e).__name__}: {e}")

    # Method 3: Try python-docx (rarely works with legacy .doc but worth trying)
    logger.info("[DOC] Trying python-docx fallback...")
    try:
        from docx import Document

        doc = Document(abs_path)
        all_text = [para.text for para in doc.paragraphs if para.text.strip()]
        if all_text:
            text = "\n".join(all_text)
            logger.info(f"[DOC] python-docx extraction successful: {len(text)} chars")
            return DocumentExtractionResult(text=text, is_ocr=False, page_count=1, confidence="medium", file_type="doc")
        else:
            logger.warning("[DOC] python-docx found no text")
    except Exception as e:
        logger.warning(f"[DOC] python-docx exception: {type(e).__name__}: {e}")

    logger.error("All DOC extraction methods failed")
    return DocumentExtractionResult(text="", confidence="low", file_type="doc")


def extract_text_from_pdf(pdf_path: str) -> DocumentExtractionResult:
    """
    Extract text from PDF file.
    First attempts native extraction, falls back to OCR if needed.
    """
    # Try native PDF extraction first
    native_text, page_count = _extract_native_text(pdf_path)

    # If we got meaningful text, return it
    if native_text and len(native_text.strip()) > 100:
        logger.info(f"Native PDF extraction successful: {len(native_text)} chars from {page_count} pages")
        return DocumentExtractionResult(
            text=native_text, is_ocr=False, page_count=page_count, confidence="high", file_type="pdf"
        )

    # Fall back to OCR
    logger.info("Native extraction yielded insufficient text, attempting OCR...")
    ocr_text, page_count = _extract_ocr_text(pdf_path)

    if ocr_text and len(ocr_text.strip()) > 50:
        logger.info(f"OCR extraction successful: {len(ocr_text)} chars from {page_count} pages")
        return DocumentExtractionResult(
            text=ocr_text, is_ocr=True, page_count=page_count, confidence="medium", file_type="pdf"
        )

    # Both methods failed
    logger.warning("Both native and OCR extraction failed or yielded minimal text")
    return DocumentExtractionResult(
        text=native_text or ocr_text or "",
        is_ocr=bool(ocr_text),
        page_count=page_count,
        confidence="low",
        file_type="pdf",
    )


def _extract_native_text(pdf_path: str) -> Tuple[str, int]:
    """Extract text using pypdf (native PDF text)."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(pdf_path)
        page_count = len(reader.pages)

        all_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                all_text.append(f"--- Page {i+1} ---\n{text}")

        return "\n\n".join(all_text), page_count
    except Exception as e:
        logger.error(f"Native PDF extraction failed: {e}")
        return "", 0


def _extract_ocr_text(pdf_path: str) -> Tuple[str, int]:
    """Extract text using OCR (for scanned PDFs)."""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        # Convert PDF pages to images
        images = convert_from_path(pdf_path, dpi=300)
        page_count = len(images)

        all_text = []
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image)
            if text:
                all_text.append(f"--- Page {i+1} (OCR) ---\n{text}")

        return "\n\n".join(all_text), page_count
    except ImportError as e:
        logger.warning(f"OCR dependencies not available: {e}")
        return "", 0
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return "", 0


def save_uploaded_document(
    file_content: bytes, filename: str, po_id: Optional[int] = None, company_id: Optional[int] = None
) -> str:
    """
    Save uploaded document (PDF or Word) to storage location.
    Returns the stored reference (local file path, or ``s3://...`` on the s3 backend).

    Persistent PO source documents must pass ``company_id`` so the remote backend can
    build a tenant-prefixed key. Callers that only need a scratch file for immediate
    text extraction (e.g. BOM import) may omit it and always get local-disk behavior.
    """
    import uuid

    from app.services.storage_service import get_local_storage, get_storage

    # Get original extension (sanitized against the supported-extension allowlist)
    original_ext = Path(filename).suffix.lower()
    if original_ext not in SUPPORTED_EXTENSIONS:
        original_ext = '.pdf'  # Default

    storage = get_storage()
    if storage.is_remote and company_id is not None:
        # Tenant-prefixed, never-user-controlled object key.
        key = f"{company_id}/purchase_orders/{po_id or 'pending'}/{uuid.uuid4()}{original_ext}"
        ref = storage.save(file_content, key=key)
        logger.info(f"Saved document to {ref}")
        return ref

    # Local layout: uploads/purchase_orders/{po_id}/{name} for PO-attached files.
    # Pending (pre-PO) files are tenant-scoped when a company is known —
    # uploads/purchase_orders/pending/{company_id}/{name} — so the serving endpoint
    # can enforce tenancy from the path alone. Callers without a company (scratch
    # extraction, e.g. BOM import) keep the legacy un-tenanted pending/ layout.
    base_dir = Path("uploads/purchase_orders")
    if po_id:
        upload_dir = base_dir / str(po_id)
    elif company_id is not None:
        upload_dir = base_dir / "pending" / str(company_id)
    else:
        upload_dir = base_dir / "pending"

    upload_dir.mkdir(parents=True, exist_ok=True)

    # Sanitize filename
    stem = Path(filename).stem
    safe_stem = "".join(c for c in stem if c.isalnum() or c in "._- ")
    safe_filename = f"{safe_stem}{original_ext}"

    file_path = upload_dir / safe_filename

    # Handle duplicate filenames
    counter = 1
    while file_path.exists():
        file_path = upload_dir / f"{safe_stem}_{counter}{original_ext}"
        counter += 1

    ref = get_local_storage().save(file_content, key=str(file_path))
    logger.info(f"Saved document to {ref}")
    return ref


def move_pdf_to_po(temp_path: str, po_id: int, company_id: Optional[int] = None) -> str:
    """Move PDF from pending to PO-specific storage after PO creation.

    On the s3 backend this is copy-to-new-key + delete-old (S3 has no move); the new
    key is server-generated under the tenant prefix. Local behavior is unchanged.

    ``temp_path`` comes from the request body, so the source ref is validated before
    any read/copy/delete: s3 refs must live in the configured bucket under this
    tenant's ``{company_id}/purchase_orders/`` prefix, and local paths must resolve
    inside ``uploads/purchase_orders`` (same guard as the PO pdf serving endpoint) —
    pending sources additionally inside ``pending/{company_id}/``.
    Violations raise ``ValueError`` (same type ``parse_s3_ref`` uses for bad refs).
    """
    import shutil
    import uuid

    from app.core.config import settings
    from app.services.storage_service import backend_for_ref, is_s3_ref, parse_s3_ref, sanitize_ext

    if is_s3_ref(temp_path):
        bucket, old_key = parse_s3_ref(temp_path)
        if bucket != settings.S3_BUCKET_NAME:
            raise ValueError(f"Source document ref is not in the configured storage bucket: {temp_path!r}")
        if company_id is not None:
            prefix = str(company_id)
        else:
            # Derive the tenant prefix from the existing (server-generated) key, but
            # only accept the canonical numeric-tenant shape.
            prefix = old_key.split("/", 1)[0]
            if not prefix.isdigit():
                raise ValueError(f"Source document ref has no tenant prefix: {temp_path!r}")
        if not old_key.startswith(f"{prefix}/purchase_orders/"):
            raise ValueError(f"Source document ref is outside the tenant's purchase_orders prefix: {temp_path!r}")
        backend = backend_for_ref(temp_path)
        if not backend.exists(temp_path):
            return temp_path
        ext = sanitize_ext(temp_path) or ".pdf"
        new_key = f"{prefix}/purchase_orders/{po_id}/{uuid.uuid4()}{ext}"
        new_ref = backend.save(backend.read_bytes(temp_path), key=new_key)
        backend.delete(temp_path)
        logger.info(f"Moved PDF from {temp_path} to {new_ref}")
        return new_ref

    # Containment guard (mirrors po_upload.py's local serving path): the source must
    # resolve inside uploads/purchase_orders before we move anything.
    allowed_root = os.path.realpath(os.path.join("uploads", "purchase_orders"))
    real_source = os.path.realpath(temp_path)
    if not real_source.startswith(allowed_root + os.sep):
        raise ValueError(f"Source document path is outside uploads/purchase_orders: {temp_path!r}")

    # The upload flow only ever produces pending/ sources locally
    # (save_uploaded_document with po_id=None), and pending sources are
    # tenant-scoped on disk (pending/{company_id}/...). Reject anything else
    # outright: a non-pending source would let a caller relocate another
    # tenant's already-attached PO document (and make it servable under their
    # own PO row). Mirrors the s3 branch's tenant-prefix check.
    segments = os.path.relpath(real_source, allowed_root).split(os.sep)
    if segments[0] != "pending":
        raise ValueError(f"Source document path is not a pending upload: {temp_path!r}")
    if company_id is not None:
        tenant_segment = str(company_id)
    else:
        # Derive the tenant segment from the path, but only accept the canonical
        # numeric-tenant shape (same rule as the s3 branch).
        tenant_segment = segments[1] if len(segments) >= 3 else ""
        if not tenant_segment.isdigit():
            raise ValueError(f"Source document path has no tenant segment: {temp_path!r}")
    if len(segments) < 3 or segments[1] != tenant_segment:
        raise ValueError(f"Source document path is outside the tenant's pending directory: {temp_path!r}")

    source = Path(temp_path)
    if not source.exists():
        return temp_path

    dest_dir = Path(f"uploads/purchase_orders/{po_id}")
    dest_dir.mkdir(parents=True, exist_ok=True)

    dest_path = dest_dir / source.name
    shutil.move(str(source), str(dest_path))

    logger.info(f"Moved PDF from {source} to {dest_path}")
    return str(dest_path)
